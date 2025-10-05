#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os, io, csv, json, zipfile, base64, sys, glob, logging, time, argparse, tempfile, shutil
from datetime import datetime
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email import encoders

# Core NLP + helpers
import spacy
from fuzzywuzzy import process
from wordfreq import top_n_list
from colorama import Fore, Style, init as color_init
import difflib

from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build


# Exports
from openpyxl import Workbook
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from docx import Document
from docx.shared import RGBColor

# Clipboard (optional)
try:
    import pyperclip
    HAS_CLIP = True
except Exception:
    HAS_CLIP = False

# Google APIs (OAuth user credentials)
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
import google.auth.exceptions
import pickle

# ================= SETTINGS (defaults; can override via CLI) =================
HIGHLIGHT = True
SEND_EMAIL = True
UPLOAD_GDOC = True
SAVE_LOCAL = True
DRY_RUN = False

OUTPUT_FILE = "clarity_output"       # base filename; timestamps appended automatically
APPEND_LOG_MD = "clarity_log.md"     # rolling log file (append mode)
ERROR_LOG = "clarity_errors.log"

EMAIL_FROM = None          # if None, Gmail API uses the authorized user ("me")
EMAIL_TO = None            # set to your email to enable email sending
DRIVE_FOLDER_ID = None     # optional Google Drive folder id for uploads

# OAuth scopes
SCOPES = [
    "https://www.googleapis.com/auth/drive.file",
    "https://www.googleapis.com/auth/gmail.send",
]

# Retries
API_RETRIES = 3
RETRY_SLEEP = 2.0
# ============================================================================

# ----- Logging -----
logging.basicConfig(filename=ERROR_LOG, filemode="a", level=logging.INFO,
                    format="%(asctime)s [%(levelname)s] %(message)s")

# ----- Console color -----
color_init(autoreset=True)

# ----- NLP -----
nlp = spacy.load("en_core_web_sm")
VALID_WORDS = set(top_n_list("en", 50000))

# Jargon and maps
JARGON_DB = {
    "asynchronous": {"simpler": "non-blocking", "category": "jargon"},
    "microservices": {"simpler": "small services", "category": "jargon"},
    "orchestration": {"simpler": "coordination", "category": "jargon"},
    "distributed tracing": {"simpler": "system tracking", "category": "jargon"},
}

SLANG_MAP = {
    "idk": "I don't know", "btw": "by the way", "imo": "in my opinion",
    "shud": "should", "gonna": "going to", "wanna": "want to",
    "cuz": "because", "tho": "though", "u": "you", "ur": "your", "pls": "please"
}

CONTRACTIONS = {
    "dont": "don't", "cant": "can't", "wont": "won't",
    "shouldnt": "shouldn't", "wouldnt": "wouldn't", "couldnt": "couldn't",
    "isnt": "isn't", "wasnt": "wasn't", "arent": "aren't", "werent": "weren't",
    "havent": "haven't", "hasnt": "hasn't", "hadnt": "hadn't",
    "doesnt": "doesn't", "didnt": "didn't", "aint": "ain't"
}

GRAMMAR_FIXES = {
    "an we": "and we",
    "an i": "and I",
    "wanna": "want to",
    "gonna": "going to",
    "cuz": "because"
}

# ----- Utility helpers -----
def ts():
    return datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

def timestamped_filename(base, ext, outdir):
    return os.path.join(outdir, f"{base}_{ts()}.{ext}")

def ensure_dir(path):
    os.makedirs(path, exist_ok=True)

def highlight(before, after):
    if HIGHLIGHT:
        return f"{Fore.RED}{before}{Style.RESET_ALL} → {Fore.GREEN}{after}{Style.RESET_ALL}"
    return f"{before} → {after}"

def google_service(api, version):
    creds = None
    if os.path.exists("token.pickle"):
        with open("token.pickle", "rb") as token:
            creds = pickle.load(token)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            try:
                creds.refresh(Request())
            except google.auth.exceptions.RefreshError:
                creds = None
        if not creds:
            flow = InstalledAppFlow.from_client_secrets_file(
    "/workspaces/portfolio/projects/clarity-coach/credentials.json", SCOPES
)
            creds = flow.run_console()
        with open("token.pickle", "wb") as token:
            pickle.dump(creds, token)
    return build(api, version, credentials=creds)

# ----- Core analyzer -----
def correct_spelling(text):
    changes = []
    categories = []  # parallel list for categories
    out = []
    for word in text.split():
        w = word.lower()
        if w in SLANG_MAP:
            changes.append(f"{word} → {SLANG_MAP[w]}"); categories.append("slang")
            out.append(SLANG_MAP[w])
        elif w in CONTRACTIONS:
            changes.append(f"{word} → {CONTRACTIONS[w]}"); categories.append("contraction")
            out.append(CONTRACTIONS[w])
        elif w in VALID_WORDS:
            out.append(word)
        else:
            match, score = process.extractOne(w, VALID_WORDS)
            if score and score > 85:
                changes.append(f"{word} → {match}"); categories.append("spelling")
                out.append(match)
            else:
                out.append(word)
    return " ".join(out), changes, categories

def smooth_grammar(text):
    changes, cats = [], []
    fixed = text
    for bad, good in GRAMMAR_FIXES.items():
        if bad in fixed:
            fixed = fixed.replace(bad, good)
            changes.append(f"{bad} → {good}")
            cats.append("grammar")
    return fixed, changes, cats

def replace_jargon(text):
    changes, cats = [], []
    final = text
    for jt, meta in JARGON_DB.items():
        if jt in final.lower():
            final = final.replace(jt, meta["simpler"])
            changes.append(f"{jt} → {meta['simpler']}")
            cats.append("jargon")
    return final, changes, cats

def sentence_cleanup(text):
    doc = nlp(text)
    sentences = []
    for sent in doc.sents:
        toks = [t.text for t in sent]
        s = " ".join(toks).replace(" ,", ",").replace(" .", ".").strip()
        if s:
            s = s[0].upper() + s[1:]
        sentences.append(s)
    return " ".join(sentences)

def analyze_text(text):
    # step 1: spelling/slang/contractions
    step1, ch1, cat1 = correct_spelling(text)
    # step 2: sentence cleanup for spacing and capitalization
    step2 = sentence_cleanup(step1)
    # step 3: grammar smoothing
    step3, ch2, cat2 = smooth_grammar(step2)
    # step 4: jargon replacement
    final, ch3, cat3 = replace_jargon(step3)

    all_changes = ch1 + ch2 + ch3
    all_cats = cat1 + cat2 + cat3

    # stats
    stats = {
        "words_before": len(text.split()),
        "words_after": len(final.split()),
        "num_changes": len(all_changes),
        "by_category": {
            "spelling": all_cats.count("spelling"),
            "slang": all_cats.count("slang"),
            "contraction": all_cats.count("contraction"),
            "grammar": all_cats.count("grammar"),
            "jargon": all_cats.count("jargon"),
        }
    }

    # prepare word-level diff for side-by-side views
    diff_pairs = compute_simple_pairs(text, final)

    return {
        "original": text,
        "polished": final,
        "changes": all_changes,
        "categories": all_cats,
        "stats": stats,
        "diff_pairs": diff_pairs  # list of (orig_chunk, polished_chunk)
    }

def compute_simple_pairs(a, b):
    # token-based diff; returns aligned pairs for HTML/PDF/DOCX side-by-side
    a_tok = a.split()
    b_tok = b.split()
    sm = difflib.SequenceMatcher(None, a_tok, b_tok)
    pairs = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            pairs.append((" ".join(a_tok[i1:i2]), " ".join(b_tok[j1:j2])))
        elif tag in ("replace", "delete", "insert"):
            pairs.append((" ".join(a_tok[i1:i2]), " ".join(b_tok[j1:j2])))
    return pairs

# ----- Console output -----
def print_console(res):
    print("=== BEFORE ==="); print(res["original"])
    print("\n=== AFTER ==="); print(res["polished"])
    print("\n=== CHANGES MADE ===")
    for i, c in enumerate(res["changes"]):
        try:
            b, a = c.split(" → ")
            print(f"- [{res['categories'][i]}] {highlight(b, a)}")
        except Exception:
            print(f"- {c}")
    print("\n=== STATS ===")
    st = res["stats"]
    print(f"Words before: {st['words_before']}")
    print(f"Words after : {st['words_after']}")
    print(f"Total changes: {st['num_changes']}")
    print("By category:", st["by_category"])

# ----- Side-by-side HTML row fragment -----
def html_diff_table_rows(diff_pairs):
    rows = []
    for left, right in diff_pairs:
        rows.append(f"<tr><td style='color:#a00'>{left}</td><td style='color:#0a0'>{right}</td></tr>")
    return "\n".join(rows)

# ----- Exports (all respect SAVE_LOCAL/DRY_RUN) -----
def save_markdown(res, outdir):
    fn = timestamped_filename(OUTPUT_FILE, "md", outdir)
    if DRY_RUN or not SAVE_LOCAL:
        return fn
    with open(fn, "w", encoding="utf-8") as f:
        f.write("# Clarity Coach Results\n\n")
        f.write("## Before\n")
        f.write(res["original"] + "\n\n")
        f.write("## After\n")
        f.write(res["polished"] + "\n\n")
        f.write("## Changes\n")
        for i, c in enumerate(res["changes"]):
            f.write(f"- [{res['categories'][i]}] {c}\n")
        f.write("\n## Stats\n")
        f.write(json.dumps(res["stats"], indent=2))
    return fn

def save_html(res, outdir):
    fn = timestamped_filename(OUTPUT_FILE, "html", outdir)
    if DRY_RUN or not SAVE_LOCAL:
        return fn
    with open(fn, "w", encoding="utf-8") as f:
        f.write("<html><body><h2>Clarity Coach Results</h2>")
        f.write("<h3>Before</h3><p>" + res["original"] + "</p>")
        f.write("<h3>After</h3><p>" + res["polished"] + "</p>")
        f.write("<h3>Side-by-side Diff</h3>")
        f.write("<table border='1' cellpadding='6' cellspacing='0'><tr><th>Original</th><th>Polished</th></tr>")
        f.write(html_diff_table_rows(res["diff_pairs"]))
        f.write("</table>")
        f.write("<h3>Changes</h3><ul>")
        for i, c in enumerate(res["changes"]):
            b, a = (c.split(" → ")+["",""])[:2]
            f.write(f"<li>[{res['categories'][i]}] <span style='color:red'>{b}</span> → <span style='color:green'>{a}</span></li>")
        f.write("</ul>")
        f.write("<h3>Stats</h3><pre>" + json.dumps(res["stats"], indent=2) + "</pre>")
        f.write("</body></html>")
    return fn

def save_txt(res, outdir):
    fn = timestamped_filename(OUTPUT_FILE, "txt", outdir)
    if DRY_RUN or not SAVE_LOCAL:
        return fn
    with open(fn, "w", encoding="utf-8") as f:
        f.write("=== BEFORE ===\n" + res["original"] + "\n\n")
        f.write("=== AFTER ===\n" + res["polished"] + "\n\n")
        f.write("=== CHANGES MADE ===\n")
        for i, c in enumerate(res["changes"]):
            f.write(f"- [{res['categories'][i]}] {c}\n")
        f.write("\n=== STATS ===\n")
        f.write(json.dumps(res["stats"], indent=2))
    return fn

def save_json(res, outdir):
    fn = timestamped_filename(OUTPUT_FILE, "json", outdir)
    if DRY_RUN or not SAVE_LOCAL:
        return fn
    with open(fn, "w", encoding="utf-8") as f:
        json.dump(res, f, indent=4, ensure_ascii=False)
    return fn

def save_csv(res, outdir):
    fn = timestamped_filename(OUTPUT_FILE, "csv", outdir)
    if DRY_RUN or not SAVE_LOCAL:
        return fn
    with open(fn, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f); w.writerow(["Category", "Original Text", "Corrected Text"])
        for i, c in enumerate(res["changes"]):
            parts = c.split(" → "); before = parts[0].strip(); after = parts[1].strip() if len(parts)>1 else ""
            w.writerow([res["categories"][i], before, after])
    return fn

def save_xlsx(res, outdir):
    fn = timestamped_filename(OUTPUT_FILE, "xlsx", outdir)
    if DRY_RUN or not SAVE_LOCAL:
        return fn
    wb = Workbook()
    ws1 = wb.active; ws1.title = "Before"; ws1["A1"] = "Original Text"; ws1["A2"] = res["original"]
    ws2 = wb.create_sheet("After"); ws2["A1"] = "Polished Text"; ws2["A2"] = res["polished"]
    ws3 = wb.create_sheet("Changes"); ws3.append(["Category", "Original Text", "Corrected Text"])
    for i, c in enumerate(res["changes"]):
        parts = c.split(" → "); before = parts[0].strip(); after = parts[1].strip() if len(parts)>1 else ""
        ws3.append([res["categories"][i], before, after])
    ws4 = wb.create_sheet("Stats"); ws4["A1"] = "Stats JSON"; ws4["A2"] = json.dumps(res["stats"])
    wb.save(fn)
    return fn

def save_pdf(res, outdir):
    fn = timestamped_filename(OUTPUT_FILE, "pdf", outdir)
    if DRY_RUN or not SAVE_LOCAL:
        return fn
    doc = SimpleDocTemplate(fn)
    styles = getSampleStyleSheet(); elems = []
    elems.append(Paragraph("Clarity Coach Results", styles["Title"])); elems.append(Spacer(1, 12))
    elems.append(Paragraph("Before:", styles["Heading2"])); elems.append(Paragraph(res["original"], styles["Normal"])); elems.append(Spacer(1, 12))
    elems.append(Paragraph("After:", styles["Heading2"])); elems.append(Paragraph(res["polished"], styles["Normal"])); elems.append(Spacer(1, 12))

    elems.append(Paragraph("Side-by-side Diff:", styles["Heading2"]))
    data = [["Original", "Polished"]]
    for left, right in res["diff_pairs"]:
        data.append([left, right])
    table = Table(data, colWidths=[250, 250])
    table.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),colors.grey), ("TEXTCOLOR",(0,0),(-1,0),colors.whitesmoke),
        ("ALIGN",(0,0),(-1,-1),"LEFT"), ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),
        ("GRID",(0,0),(-1,-1),0.5,colors.black)
    ]))
    elems.append(table); elems.append(Spacer(1, 12))

    elems.append(Paragraph("Changes:", styles["Heading2"]))
    data2 = [["Category", "Original", "Corrected"]]
    for i, c in enumerate(res["changes"]):
        parts = c.split(" → "); b = parts[0]; a = parts[1] if len(parts)>1 else ""
        data2.append([res["categories"][i], b, a])
    table2 = Table(data2, colWidths=[80, 210, 210])
    table2.setStyle(TableStyle([("GRID",(0,0),(-1,-1),0.5,colors.black)]))
    elems.append(table2); elems.append(Spacer(1, 12))

    elems.append(Paragraph("Stats:", styles["Heading2"]))
    elems.append(Paragraph(json.dumps(res["stats"]), styles["Code"]))

    doc.build(elems)
    return fn

def save_docx(res, outdir):
    fn = timestamped_filename(OUTPUT_FILE, "docx", outdir)
    if DRY_RUN or not SAVE_LOCAL:
        return fn
    d = Document()
    d.add_heading("Clarity Coach Results", 0)
    d.add_heading("Before", 1); d.add_paragraph(res["original"])
    d.add_heading("After", 1); d.add_paragraph(res["polished"])

    d.add_heading("Side-by-side Diff", 1)
    for left, right in res["diff_pairs"]:
        p = d.add_paragraph()
        runL = p.add_run(left + "  "); runL.font.color.rgb = RGBColor(0xAA, 0x00, 0x00)
        runR = p.add_run(right);      runR.font.color.rgb = RGBColor(0x00, 0xAA, 0x00)

    d.add_heading("Changes", 1)
    for i, c in enumerate(res["changes"]):
        p = d.add_paragraph(f"[{res['categories'][i]}] ")
        parts = c.split(" → ")
        runL = p.add_run(parts[0] + " "); runL.font.color.rgb = RGBColor(0xAA, 0x00, 0x00)
        if len(parts) > 1:
            p.add_run("→ ")
            runR = p.add_run(parts[1]); runR.font.color.rgb = RGBColor(0x00, 0xAA, 0x00)

    d.add_heading("Stats", 1)
    d.add_paragraph(json.dumps(res["stats"], indent=2))
    d.save(fn)
    return fn

def save_gdoc(res, tmpdir):
    # create temp txt to convert
    temp = timestamped_filename(OUTPUT_FILE, "txt", tmpdir)
    with open(temp, "w", encoding="utf-8") as f:
        f.write("=== BEFORE ===\n" + res["original"] + "\n\n")
        f.write("=== AFTER ===\n" + res["polished"] + "\n\n")
        f.write("=== CHANGES ===\n")
        for i, c in enumerate(res["changes"]):
            f.write(f"- [{res['categories'][i]}] {c}\n")
        f.write("\n=== STATS ===\n")
        f.write(json.dumps(res["stats"], indent=2))

    if DRY_RUN:
        return temp, None

    drive = google_service("drive", "v3")
    meta = {"name": os.path.splitext(os.path.basename(temp))[0], "mimeType": "application/vnd.google-apps.document"}
    if DRIVE_FOLDER_ID: meta["parents"] = [DRIVE_FOLDER_ID]
    file_id = None; link = None
    for attempt in range(API_RETRIES):
        try:
            media = MediaFileUpload(temp, mimetype="text/plain")
            file = drive.files().create(body=meta, media_body=media, fields="id,webViewLink").execute()
            file_id = file.get("id"); link = file.get("webViewLink")
            break
        except Exception as e:
            logging.error(f"Drive upload failed: {e}")
            time.sleep(RETRY_SLEEP)
    return temp, link

# ----- Zip and email -----
def zip_files(paths, outdir):
    fn = timestamped_filename(OUTPUT_FILE, "zip", outdir)
    if DRY_RUN:
        return fn
    with zipfile.ZipFile(fn, "w", zipfile.ZIP_DEFLATED) as z:
        for p in paths:
            if p and os.path.exists(p):
                z.write(p, arcname=os.path.basename(p))
    return fn

def send_email_with_zip(zip_path, subject="Clarity Coach Results", body="See attached results."):
    if not zip_path or not os.path.exists(zip_path) or not SEND_EMAIL:
        return False
    if EMAIL_TO is None:
        logging.info("EMAIL_TO not set; skipping email.")
        return False

    gmail = google_service("gmail", "v1")
    sender = EMAIL_FROM if EMAIL_FROM else "me"

    message = MIMEMultipart()
    message["to"] = EMAIL_TO
    message["from"] = sender
    message["subject"] = subject
    message.attach(MIMEText(body, "plain"))

    with open(zip_path, "rb") as f:
        part = MIMEBase("application", "zip")
        part.set_payload(f.read())
    encoders.encode_base64(part)
    part.add_header("Content-Disposition", f'attachment; filename="{os.path.basename(zip_path)}"')
    message.attach(part)

    raw = base64.urlsafe_b64encode(message.as_bytes()).decode()

    for attempt in range(API_RETRIES):
        try:
            gmail.users().messages().send(userId="me", body={"raw": raw}).execute()
            return True
        except Exception as e:
            logging.error(f"Gmail send failed: {e}")
            time.sleep(RETRY_SLEEP)
    return False

# ----- Append log -----
def append_log(res):
    try:
        with open(APPEND_LOG_MD, "a", encoding="utf-8") as f:
            f.write(f"\n\n# Run {ts()}\n")
            f.write("## Before\n" + res["original"] + "\n\n")
            f.write("## After\n" + res["polished"] + "\n\n")
            f.write("## Changes\n")
            for i, c in enumerate(res["changes"]):
                f.write(f"- [{res['categories'][i]}] {c}\n")
            f.write("\n## Stats\n" + json.dumps(res["stats"], indent=2) + "\n")
    except Exception as e:
        logging.error(f"Append log failed: {e}")

# ----- Clipboard -----
def copy_to_clipboard(text):
    if HAS_CLIP:
        try:
            pyperclip.copy(text)
        except Exception as e:
            logging.error(f"Clipboard copy failed: {e}")

# ----- Orchestrator for one text -----
def process_text(text, outroot):
    res = analyze_text(text)

    # console
    print_console(res)

    # clipboard
    copy_to_clipboard(res["polished"])

    # outputs
    produced = []
    if SAVE_LOCAL:
        produced.append(save_markdown(res, outroot))
        produced.append(save_html(res, outroot))
        produced.append(save_txt(res, outroot))
        produced.append(save_json(res, outroot))
        produced.append(save_csv(res, outroot))
        produced.append(save_xlsx(res, outroot))
        produced.append(save_pdf(res, outroot))
        produced.append(save_docx(res, outroot))

    # Google Doc
    gdoc_temp = None
    gdoc_link = None
    if UPLOAD_GDOC:
        gdoc_temp, gdoc_link = save_gdoc(res, outroot)
        if gdoc_temp: produced.append(gdoc_temp)
        if gdoc_link: print(f"Google Doc: {gdoc_link}")

    # Zip + email
    zip_path = zip_files(produced, outroot) if (SAVE_LOCAL or UPLOAD_GDOC) else None
    if SEND_EMAIL and EMAIL_TO:
        ok = send_email_with_zip(zip_path)
        print("Email sent." if ok else "Email skipped or failed.")

    # append log
    append_log(res)

    return res

# ----- CLI -----
def main():
    parser = argparse.ArgumentParser(description="Clarity Coach: correct spelling, grammar, jargon; export everywhere.")
    parser.add_argument("text", nargs="?", help="Text to process (wrap in quotes). If omitted, use --folder.")
    parser.add_argument("--folder", help="Process all .txt files in folder (batch mode).")
    parser.add_argument("--no-email", action="store_true", help="Disable email sending.")
    parser.add_argument("--no-gdoc", action="store_true", help="Disable Google Docs upload.")
    parser.add_argument("--no-save", action="store_true", help="Do not save local files.")
    parser.add_argument("--dry-run", action="store_true", help="Analyze only; no save/upload/email.")
    parser.add_argument("--email-to", help="Override EMAIL_TO.")
    parser.add_argument("--email-from", help="Override EMAIL_FROM.")
    parser.add_argument("--drive-folder", help="Google Drive folder ID.")
    parser.add_argument("--outdir", default="clarity_outputs", help="Output directory.")
    args = parser.parse_args()


    global SEND_EMAIL, UPLOAD_GDOC, SAVE_LOCAL, DRY_RUN, EMAIL_TO, EMAIL_FROM, DRIVE_FOLDER_ID
    if args.no_email: SEND_EMAIL = False
    if args.no_gdoc: UPLOAD_GDOC = False
    if args.no_save: SAVE_LOCAL = False
    if args.dry_run: DRY_RUN = True
    if args.email_to: EMAIL_TO = args.email_to
    if args.email_from: EMAIL_FROM = args.email_from
    if args.drive_folder: DRIVE_FOLDER_ID = args.drive_folder

    outdir = args.outdir
    if SAVE_LOCAL and not DRY_RUN:
        ensure_dir(outdir)
    else:
        ensure_dir(outdir)  # still host temp artifacts

    try:
        if args.folder:
            # batch mode: each .txt file processed
            files = sorted(glob.glob(os.path.join(args.folder, "*.txt")))
            if not files:
                print("No .txt files found in folder.")
                return
            for fp in files:
                # Skip previously generated output files to avoid recursion
                if "clarity_output" in os.path.basename(fp):
                    continue

                with open(fp, "r", encoding="utf-8") as f:
                    text = f.read().strip()
                print(f"\n=== Processing: {fp} ===")
                process_text(text, outdir)

        else:
            # single text mode
            if not args.text:
                print("Provide text or use --folder.")
                return
            process_text(args.text, outdir)
    except Exception as e:
        logging.error(f"Fatal error: {e}")
        print("An error occurred. See clarity_errors.log for details.")

if __name__ == "__main__":
    main()
